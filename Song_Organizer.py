import os
from tinytag import TinyTag
from docx import Document
from docx.shared import Pt
import tkinter as tk
from tkinter import *
from tkinter import ttk
from docx.shared import Inches

# Tkinter Window

window = Tk()
window.title("Search Settings")
window.minsize(width=300, height=300)
window.geometry('600x600+0+0')

cd1 = StringVar()
cd1.set('R 01')
cd1_label=Label(window, text="Enter the name of cd 1: ").place(x = 5, y = 10)
cd1_entry = Entry(window,width=20,textvariable=cd1).place(x=180, y=10)

cd2 = StringVar()
cd2.set('R 02')
cd2_label=Label(window, text="Enter the name of cd 2: ").place(x = 5, y = 40)
cd2_entry = Entry(window,width=20,textvariable=cd2).place(x=180, y=40)

cd3 = StringVar()
cd3.set('R 03')
cd3_label=Label(window, text="Enter the name of cd 3: ").place(x = 5, y = 70)
cd3_entry = Entry(window,width=20,textvariable=cd3).place(x=180, y=70)

cd4 = StringVar()
cd4.set('R 04')
cd4_label=Label(window, text="Enter the name of cd 4: ").place(x = 5, y = 100)
cd4_entry = Entry(window,width=20,textvariable=cd4).place(x=180, y=100)

font_size = IntVar()
font_size.set('8')
font_label=Label(window, text="Enter the font size you want: ").place(x = 5, y = 130)
font_entry = Entry(window,width=10,textvariable=font_size).place(x=200, y=130)

file_name = StringVar()
file_name.set('KenSongs')
file_label=Label(window, text="Enter the song file name: ").place(x = 5, y = 160)
file_entry = Entry(window,width=20,textvariable=file_name).place(x=280, y=160)

output_file = StringVar()
output_file.set('C:/Users/gnork/Music/Music/')

output_label=Label(window, text="Enter where the song file goes: ").place(x = 5, y = 190)
output_entry = Entry(window,width=30,textvariable=output_file).place(x=280, y=190)

rootdir = StringVar()
rootdir.set('C:/Users/gnork/Music/Music/')

rootdir_label=Label(window, text="Enter where in your music folder is: ").place(x = 5, y = 220)
rootdir_entry = Entry(window,width=30,textvariable=rootdir).place(x=280, y=220)

def run_everything():
    cd_list=[]
    cd_list.append(cd1.get())
    cd_list.append(cd2.get())
    cd_list.append(cd3.get())
    cd_list.append(cd4.get())

    document = Document()
    style = document.styles['Normal']
    sections = document.sections
    for section in sections:
        section.left_margin = Inches(.75)
        section.right_margin = Inches(.25)
        section.top_margin = Inches(.5)
        section.bottom_margin = Inches(.25)
    font = style.font
    font.name = 'Arial'
    font.size = Pt(font_size.get())

    count=0
    cd_stuff1=[]
    cd_stuff2=[]
    song1=[]
    song2=[]
    artist1=[]
    artist2=[]
    for subdir, dirs, files in sorted(os.walk(rootdir.get())):
        if set(cd_list).issubset(dirs):
            dirs=cd_list.copy()
        if os.path.split(subdir)[1] in cd_list:
            if len(dirs)>0:
                continue
            if os.path.split(subdir)[1]!="":
                count+=1
                print("Loading cd", os.path.split(subdir)[1])

                if count%2==0:
                    cd_stuff2.append("CD:" + os.path.split(subdir)[1])
                    #cd_stuff2.append("Song" + " | " + "Artist")
                else:
                    cd_stuff1.append("CD:" + os.path.split(subdir)[1])
                    #cd_stuff1.append("Song" + " | " + "Artist")
            for file in files:
                if file[-4:]==".wav":
                    tag= TinyTag.get(subdir+"/"+file)
                    print("Reading:", file[:-4], "by", tag.artist)
                    if count%2==0:
                        song2.append(file[:-4])
                        artist2.append(tag.artist)
                    else:
                        song1.append(file[:-4])
                        artist1.append(tag.artist)

            if count%2==0:
                table = document.add_table(rows=1, cols=4)
                table.style = 'TableGrid'

                hdr_cells = table.rows[0].cells
                h1=hdr_cells[0].paragraphs[0].add_run(cd_stuff1[0])
                h1.bold=True
                h2=hdr_cells[2].paragraphs[0].add_run(cd_stuff2[0])
                h2.bold=True

                max_range=max(len(song1),len(song2))
                for i in range(0,max_range):
                    row_cells = table.add_row().cells
                    try:
                        row_cells[0].text = str(song1[i])
                        row_cells[1].text = str(artist1[i])
                    except:
                        pass
                    try:
                        row_cells[2].text = str(song2[i])
                        row_cells[3].text = str(artist2[i])
                    except:
                        pass

                document.add_paragraph(' ')

                cd_stuff1=[]
                cd_stuff2=[]
                artist1=[]
                artist2=[]
                song1=[]
                song2=[]
            if count==4:
                count=0

    document.save(output_file.get()+file_name.get()+'.docx')

    print(" ")
    print("Wrote song list file to this location:", output_file.get())
    print("Song list file is called ", file_name.get() +".docx")

final_button =Button(window, text="Run program", command=run_everything)
final_button.place(x = 5, y = 390)

window.mainloop()

# End of Tkinter window
